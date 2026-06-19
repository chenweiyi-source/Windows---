from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Cm(3.0)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

def set_font(run, size=Pt(12), bold=False):
    run.font.name = '宋体'
    run.font.size = size
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._element.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')

def add_para(doc, text, bold=False, indent=True, align=WD_ALIGN_PARAGRAPH.LEFT, size=Pt(12)):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    set_font(run, size, bold)
    return p

def add_code_line(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p

def add_img(doc, path, w=14):
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run()
        run.add_picture(path, width=Cm(w))
    except:
        add_para(doc, '[此处插入界面效果图]')

# ==================== CONTENT ====================

add_para(doc, 'Windows程序设计课程作业', bold=True, indent=False,
         align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(16))

# ===== 1. 课程综述 =====
add_para(doc, '1. 课程综述', bold=True, indent=False, size=Pt(14))

add_para(doc, '（一）课程总结（20分）', bold=True)

add_para(doc,
    '本学期通过《Windows程序设计》课程的学习，系统掌握了使用C#进行Windows桌面应用程序开发的核心技术体系。'
    '在C#语法与面向对象编程方面，深入理解了封装、继承、多态、抽象四大特性，以及泛型、委托、事件等高级特性，'
    '通过"学生信息管理系统"实验项目的类设计实践，掌握了面向对象分析与设计的实际应用。'
    '在WinForm控件与事件驱动编程方面，系统学习了Label、Button、TextBox、DataGridView、ListView等常用控件的'
    '属性配置、事件绑定和方法调用，通过"记事本"实验项目深入体会了事件驱动模型的工作机制。'
    '文件操作方面掌握了System.IO命名空间下File、Directory、FileStream、StreamReader等类的使用，'
    '实现了文件的创建、读写、复制、移动等操作，并在异常处理方面采用了try-catch-finally结构确保程序健壮性。'
    '多媒体处理方面学习了axWindowsMediaPlayer控件的使用，包括play、pause、stop等播放控制接口调用，'
    '以及音量调节、进度控制、状态事件监听等技术。'
    '数据库操作方面掌握了ADO.NET的基本流程：连接建立（SqlConnection）、命令执行（SqlCommand）、'
    '数据读取（SqlDataReader）和结果展示（DataGridView绑定）。'
    'GDI+绘图方面学习了Graphics对象的获取、Pen和Brush的使用、基本图形绘制等，在"简易画图板"实验中进行了实践。'
    '以上各项技术在期末项目开发中得到了综合运用和进一步巩固。')

add_para(doc, '（二）技术深入与功能实现（10分）', bold=True)

add_para(doc,
    '在课程所学技术中，我选择多媒体控制技术作为深入研究对象。选择理由如下：第一，多媒体控制涉及WinForm控件编程、'
    '事件驱动、COM组件调用等多个知识点的综合运用，能够全面检验课程所学。第二，该技术具有较高的实用价值，'
    '媒体播放是桌面应用最常见的功能之一。第三，该技术具有良好的扩展性，可在基础功能之上扩展播放列表管理、'
    '歌词显示等功能。')

add_para(doc,
    '功能设计思路：设计一个媒体播放控制函数ControlPlayback，通过MediaState枚举管理播放状态，'
    '利用switch语句根据目标状态执行对应操作。该设计采用状态机思想，确保播放、暂停、停止之间的状态转换合法且可控。'
    '同时对文件路径为空、文件不存在、操作不允许等异常情况进行了完整检查和处理。')

add_para(doc, '核心实现代码：')

code = [
    'public enum MediaState { Stopped, Playing, Paused }',
    '',
    'public bool ControlPlayback(MediaState target, string path = "") {',
    '    try {',
    '        switch (target) {',
    '            case MediaState.Playing:',
    '                if (string.IsNullOrEmpty(path))',
    '                    throw new ArgumentException("播放路径不能为空");',
    '                if (!File.Exists(path))',
    '                    throw new FileNotFoundException("找不到媒体文件", path);',
    '                player.URL = path;',
    '                player.Ctlcontrols.play();',
    '                currentState = MediaState.Playing; break;',
    '            case MediaState.Paused:',
    '                if (currentState != MediaState.Playing)',
    '                    throw new InvalidOperationException("当前状态不可暂停");',
    '                player.Ctlcontrols.pause();',
    '                currentState = MediaState.Paused; break;',
    '            case MediaState.Stopped:',
    '                player.Ctlcontrols.stop();',
    '                currentState = MediaState.Stopped; break;',
    '        } return true;',
    '    } catch (Exception ex) {',
    '        MessageBox.Show($"操作失败：{ex.Message}"); return false;',
    '    }',
    '}'
]
for line in code:
    add_code_line(doc, line)

add_para(doc,
    '该函数的技术要点：一是通过状态枚举保证操作合法性，只有播放状态下允许暂停，体现了严谨的状态管理。'
    '二是对多种异常情况进行完整检查并给出友好提示。三是将控制逻辑封装在独立方法中，遵循单一职责原则。'
    '通过该功能的实现，加深了对多媒体控制技术、状态机设计和异常处理最佳实践的理解。')

# ===== 2. UI设计及说明 =====
add_para(doc, '2. UI设计及说明（10分）', bold=True, indent=False, size=Pt(14))

add_para(doc,
    '2.1 设计参考来源', bold=True)

add_para(doc,
    '本次UI设计参考了以下资源：一是B站WinForm高颜值界面设计教程，学习了GDI+自绘控件、布局优化、配色方案等技巧。'
    '二是微软官方Windows Forms控件文档，系统查阅了各类控件的属性、事件和方法。'
    '三是参考了现代音乐播放器（如Spotify、网易云音乐）的界面布局和视觉风格。')

add_para(doc,
    '2.2 设计理念与配色方案', bold=True)

add_para(doc,
    '设计定位为一款现代风格的音乐播放器。整体采用深色主题（背景色#121223），搭配珊瑚红（#FF6B6B）作为主色调、'
    '金色（#FFD764）作为强调色、青色（#45B7D1）作为辅助色，营造出优雅、沉浸的视觉氛围。'
    '界面布局采用经典的三段式结构：顶部标题栏、中部内容区、底部控制栏。中部内容区又分为左侧导航栏、'
    '中间歌曲列表和右侧正在播放面板，信息层级清晰，用户体验流畅。')

add_para(doc,
    '核心设计亮点包括：一是使用LinearGradientBrush绘制标题栏和控制栏的渐变背景，替代纯色填充，提升视觉层次感。'
    '二是在专辑封面区域使用GDI+绘制CD光盘装饰图案，包含渐变填充、外圈光晕和内圈文字，增添了界面的精致度。'
    '三是对ListView进行OwnerDraw自定义绘制，实现隔行变色、选中高亮、自定义字体颜色等效果，'
    '使列表展示更加美观。四是为所有交互元素添加悬停变色效果，提供即时操作反馈。')

add_img(doc, r'D:\C#\期末作业\MusicPlayer\screenshot.png', 14)

add_para(doc,
    '2.3 控件使用说明', bold=True)

add_para(doc,
    '本界面综合使用了多种Windows Forms控件，具体如下：'
    '（1）Panel控件：作为布局容器，将窗口划分为标题栏、左侧导航、主内容区、右侧面板、控制栏五个区域。'
    '通过设置BackColor实现视觉分区，并在Panel的Paint事件中绘制渐变背景和分隔线。'
    '（2）Button控件：用作导航菜单项和播放控制按钮。通过FlatStyle.Flat去除默认立体边框，自定义BackColor和ForeColor，'
    '绑定MouseEnter/MouseLeave事件实现悬停变色效果，Click事件实现功能切换。'
    '（3）ListView控件：以OwnerDraw模式展示歌曲列表。自定义DrawItem事件处理实现隔行交替背景色、'
    '选中行金色高亮、各列独立字体颜色等效果，相比默认样式更加美观。'
    '（4）Label控件：用于显示标题、时间、播放信息等文本。通过Font、ForeColor、AutoSize等属性精细控制样式。'
    '（5）TextBox控件：用于搜索输入，设置深色背景和浅色文字融入整体主题。'
    '（6）TrackBar控件：用作播放进度条和音量滑块。通过TickStyle.None隐藏刻度线，自定义ForeColor匹配主题色。'
    '（7）Timer控件（未直接使用在UI中，但可在后续扩展中实现进度更新功能）。')

add_para(doc,
    '2.4 关键技术要点', bold=True)

add_para(doc,
    '一是纯代码布局。所有控件均在SetupUI方法中通过C#代码精确创建和配置，未依赖WinForm设计器拖拽，'
    '体现了对控件属性、事件、方法的全面掌握。'
    '二是GDI+自绘技术。在Panel的Paint事件和ListView的DrawItem事件中使用Graphics对象进行自定义绘制，'
    '实现了渐变背景、圆形装饰图案、自定义列表样式等效果，提升了界面视觉表现力。'
    '三是事件驱动编程。为按钮绑定多种鼠标事件实现丰富的交互反馈，充分体现了WinForm事件驱动的编程思想。'
    '四是颜色管理。统一管理所有颜色常量，确保界面配色协调一致，体现了良好的代码组织习惯。')

add_para(doc,
    '2.5 总结', bold=True)

add_para(doc,
    '本次UI设计综合运用了九种Windows Forms控件，采用深色主题的现代设计风格，以纯代码方式实现了布局精致、'
    '交互流畅的音乐播放器界面。设计过程中重点体现了对控件知识、GDI+绘图、事件驱动编程等课程核心内容的掌握。')

# ========== Save ==========
out = r'D:\C#\期末作业\课程综述_第一二部分.docx'
doc.save(out)

total = sum(len(p.text) for p in doc.paragraphs if p.text.strip())
print(f'保存成功。总字符数: {total}')
