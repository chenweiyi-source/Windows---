from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21.0)
sec.page_height = Cm(29.7)
sec.left_margin = Cm(3.0)
sec.right_margin = Cm(2.5)
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.5)

def sf(run, size=Pt(12), bold=False):
    run.font.name = '宋体'
    run.font.size = size
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._element.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')

def ap(doc, text, bold=False, indent=True, align=WD_ALIGN_PARAGRAPH.LEFT, size=Pt(12)):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if indent: p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    sf(run, size, bold)
    return p

def ac(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p

def aimg(doc, path, w=13):
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run()
        run.add_picture(path, width=Cm(w))
    except:
        ap(doc, '[此处插入截图]')

# ======================= 标题 =======================
ap(doc, 'Windows程序设计课程作业', bold=True, indent=False,
   align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(16))

# ======================= 1. 课程综述 =======================
ap(doc, '1. 课程综述', bold=True, indent=False, size=Pt(14))
ap(doc, '1.1 课程总结（20分）', bold=True)

ap(doc,
    '本学期系统学习了使用C#进行Windows桌面应用开发的核心技术。C#面向对象方面掌握了封装、继承、多态、抽象四大特性'
    '及泛型、委托、事件等高级特性，在音乐播放器和DailyNote项目中通过设计Model类和封装业务逻辑进行了实践。'
    'WinForm控件方面学习了Label、Button、TextBox、DataGridView、ListView、ListBox、Panel、'
    'DateTimePicker、ComboBox、Timer等控件的属性、事件和方法。事件驱动编程是WinForm核心机制，在ZooMonitor'
    '中为DataGridView的SelectionChanged事件绑定联动更新，在DailyNote中通过DrawItem实现自定义绘制。')

ap(doc,
    '多媒体处理方面使用axWindowsMediaPlayer控件实现音视频播放控制，包括Ctlcontrols.play()/pause()/stop()、'
    'settings.volume和currentPosition等接口，在音乐播放器中实现了对OGG/MP3文件的播放，在流媒体播放器中'
    '实践了网络直播流播放。文件操作方面学习了System.IO命名空间下的File、Directory、FileStream等类，'
    '在DailyNote中使用File.ReadAllText/WriteAllText结合System.Text.Json实现数据持久化。'
    'GDI+绘图方面学习了Graphics对象、Pen/Brush使用及基本图形绘制，在图片处理中实践了裁剪、打马赛克、'
    '颜色变换，在ZooMonitor中使用Graphics绘制健康指标柱状图。多线程方面学习了Thread和Task的使用。'
    '数据库方面掌握了ADO.NET基本流程。异常处理方面在文件读写、数据解析、播放控制等环节使用了try-catch机制。')

ap(doc, '1.2 技术深入与功能实现（10分）', bold=True)

ap(doc,
    '在课程所学的多项技术中，我选择多媒体控制技术作为深入研究和功能实现的对象。做出这一选择的理由如下：'
    '第一，多媒体控制技术综合运用了WinForm控件编程、事件驱动机制、COM组件调用、状态机设计等多个课程知识点，'
    '是最能全面检验和体现课程所学掌握程度的技术领域。第二，该技术具有极高的实用价值，无论是音乐播放器、'
    '流媒体播放器还是视频监控系统，多媒体播放控制都是核心功能模块。第三，该技术具有良好的功能扩展性，'
    '在基础的播放、暂停、停止功能之上，可以进一步扩展播放列表管理、歌词同步显示、音效均衡器调节等高级功能。')

ap(doc,
    '基于上述技术选择，我设计实现了一个通用的媒体播放控制函数。功能设计思路如下：首先，通过MediaState枚举'
    '定义三种播放状态——Playing（播放中）、Paused（已暂停）、Stopped（已停止），利用状态机思想确保各种'
    '播放操作之间的状态转换合法可控。其次，通过switch语句根据传入的目标状态参数执行对应的播放控制操作。'
    '最后，在每步操作中添加完善的异常处理机制，针对文件路径为空、文件不存在、操作不允许等异常场景分别给出'
    '明确的错误提示。该函数封装在独立的MediaController类中，与界面逻辑分离，体现了单一职责的设计原则。')

ap(doc, '核心实现代码如下：')
ac(doc, 'public enum MediaState { Stopped, Playing, Paused }')
ac(doc, '')
ac(doc, 'public bool ControlPlayback(MediaState target, string path = "")')
ac(doc, '{')
ac(doc, '    try {')
ac(doc, '        switch (target) {')
ac(doc, '            case MediaState.Playing:')
ac(doc, '                if (string.IsNullOrEmpty(path))')
ac(doc, '                    throw new ArgumentException("播放路径不能为空");')
ac(doc, '                if (!File.Exists(path))')
ac(doc, '                    throw new FileNotFoundException("找不到文件");')
ac(doc, '                player.URL = path;')
ac(doc, '                player.Ctlcontrols.play();')
ac(doc, '                state = MediaState.Playing; break;')
ac(doc, '            case MediaState.Paused:')
ac(doc, '                if (state != MediaState.Playing)')
ac(doc, '                    throw new InvalidOperationException("当前无法暂停");')
ac(doc, '                player.Ctlcontrols.pause();')
ac(doc, '                state = MediaState.Paused; break;')
ac(doc, '            case MediaState.Stopped:')
ac(doc, '                player.Ctlcontrols.stop();')
ac(doc, '                state = MediaState.Stopped; break;')
ac(doc, '        }')
ac(doc, '        return true;')
ac(doc, '    }')
ac(doc, '    catch (Exception ex)')
ac(doc, '    { MessageBox.Show($"操作失败：{ex.Message}"); return false; }')
ac(doc, '}')

ap(doc,
    '该函数的设计体现了三点核心技术要点：一是通过状态枚举严格保证操作合法性，只有在Playing状态下才允许执行'
    '暂停操作，避免了在已暂停或已停止状态下调用暂停方法导致的运行时错误，体现了严谨的状态机设计思想。'
    '二是对文件路径为空、文件不存在等常见异常情况进行了预防性检查并给出中文友好提示，提升了程序的健壮性和'
    '用户体验。三是将播放控制逻辑封装在独立的方法中并进一步封装在MediaController类中，与界面代码解耦，'
    '遵循了面向对象的单一职责原则，便于后续维护和功能扩展。通过该功能的完整实现，不仅巩固了对多媒体控制'
    '技术的掌握，还加深了对状态机设计模式和异常处理最佳实践的理解。')

# ======================= 2. UI设计及说明 =======================
ap(doc, '2. UI设计及说明（10分）', bold=True, indent=False, size=Pt(14))

ap(doc,
    '本次UI设计的对象是ZooMonitor动物园动物健康监测系统，这是一个面向动物园管理人员的可视化数据监测工具。'
    '设计定位为管理型仪表盘界面，功能涵盖动物列表概览、个体健康数据可视化、观察记录管理、实时警报和活动日志'
    '追踪等。设计过程中参考了B站WinForm界面设计教程中的布局技巧和配色方法，以及微软官方控件文档中关于'
    'DataGridView、ListBox等控件的高级用法。同时参考了主流数据监控仪表盘的界面风格，采用卡片式数据展示和'
    '图表化指标呈现。')

ap(doc,
    '整体采用自然清新的配色方案：深绿色（RGB:45,85,65）作为顶部标题栏主色，传达自然和生命的寓意；'
    '主体背景为暖白色（RGB:240,245,235），四个统计卡片分别使用蓝色（总数）、绿色（健康）、黄色（观察中）、'
    '红色（需关注）作为标识色，让用户一目了然地掌握整体状况。这种配色借鉴了交通信号灯的视觉逻辑，'
    '绿色代表正常、黄色代表警戒、红色代表异常，符合用户的直觉认知习惯。')

ap(doc,
    '界面布局采用经典的后台管理系统结构：顶部为标题栏（显示系统名称、副标题和实时时间），接着是四个统计概览卡片'
    '（横向排列展示动物总数、健康数量、观察中数量、需关注数量），中部采用左-中-右三栏布局（左侧为动物列表含搜索功能、'
    '中间为健康详情含五项指标柱状图和观察记录编辑区、右侧为活动日志时间线），底部为状态提示栏。')

ap(doc,
    '界面截图如下：')
aimg(doc, r'D:\C#\期末作业\ZooMonitor\screenshot.png', 14)

ap(doc,
    '本界面综合使用了以下Windows Forms控件，体现了对控件知识的掌握：'
    '（1）Panel控件：作为核心布局容器，将界面划分为标题栏、统计卡片行、左侧动物列表区、中间健康详情区、'
    '右侧活动日志区等五个功能区域。每个Panel通过BackColor属性设置不同的背景色实现视觉分区，'
    '并在Paint事件中绘制圆角矩形和渐变背景效果，提升了界面的精致度。')

ap(doc,
    '（2）DataGridView控件：用于展示动物列表，包含姓名、种类、状态三列。通过自定义ColumnHeadersDefaultCellStyle'
    '设置深绿色表头样式，通过RowsDefaultCellStyle设置白色行样式和绿色选中高亮，通过SelectionChanged事件'
    '绑定动物详情更新逻辑。特别对状态列的ForeColor进行了动态设置，健康为绿色、观察中为黄色、需关注为红色，'
    '实现了条件格式化的视觉效果。')

ap(doc,
    '（3）ListBox控件：在右侧活动日志区域使用ListBox展示按时间倒序排列的饲养记录，每项包含时间戳和活动描述，'
    '使用自定义BackColor融入整体浅色主题。在警报区域也使用了ListBox展示当前需要关注的异常情况，'
    '红色文字配合浅红背景突出警示效果。')

ap(doc,
    '（4）Label控件：用于显示标题、统计数字、区域标题等文本信息。统计卡片中的数字使用28pt大号加粗字体突出显示，'
    '区域标题使用10pt加粗字体配合深色文字，层次分明。')

ap(doc,
    '（5）TextBox控件：用于观察记录编辑和搜索输入。观察记录编辑框设置为Multiline模式支持多行文本输入，'
    '方便饲养员记录详细观察信息。搜索框支持按动物名称快速筛选。')

ap(doc,
    '（6）Button控件：用于"Update Notes"更新观察记录、"Add Record"添加活动记录、"Export Report"导出报告等操作。'
    '通过FlatStyle.Flat去掉默认立体边框，配合圆角矩形和不同颜色区分功能类型。')

ap(doc,
    '（7）Timer控件：用于顶部状态栏的时间更新，每30秒刷新一次显示当前时间。')

ap(doc,
    '关键技术要点：一是纯代码布局，所有控件通过C#代码精确控制位置、大小、颜色和字体，体现了对控件属性的全面掌握。'
    '二是GDI+绘图在统计卡片中的应用，使用GraphicsPath绘制圆角矩形替代默认矩形，提升视觉精致度。'
    '三是在健康详情区域的Panel中通过Paint事件使用Graphics对象绘制五项健康指标的柱状图可视化，'
    '用不同颜色的填充矩形直观表示心率、体温、活动量、食欲、体重五个指标的健康程度。'
    '四是条件格式化的应用，根据数据状态动态设置文字颜色（健康绿色、警戒黄色、危险红色），'
    '使用户能够快速识别异常数据。五是事件驱动编程的灵活运用，为DataGridView的SelectionChanged事件'
    '绑定联动更新逻辑，选中不同动物时自动切换健康详情。')

ap(doc,
    '本次UI设计以控件知识的掌握为核心目标，综合运用了七种Windows Forms控件，采用管理型仪表盘的界面结构，'
    '实现了功能清晰、视觉舒适的动物园健康监测系统界面。设计过程中重点体现了DataGridView的高级配置、'
    'GDI+数据可视化、条件格式化等关键技术，为面向管理信息系统的WinForm界面设计积累了实践经验。')

# ======================= 3. AI协同项目实践 =======================
ap(doc, '3. 应用开发——AI协同项目实践（30分）', bold=True, indent=False, size=Pt(14))
ap(doc, '3.1 功能模块说明与核心代码', bold=True)

ap(doc,
    'DailyNote日记本应用围绕日记条目的增删改查核心需求，划分为以下五个功能模块：')

ap(doc,
    '（1）数据模型模块。定义DiaryEntry日记条目类，包含Id（唯一标识）、Title（标题）、Content（正文内容）、'
    'Mood（心情标签）、Date（日记日期）、CreatedAt（创建时间）六个字段。Id使用Guid生成并截取前8位作为短标识符，'
    '确保每条日记具有全局唯一性。Mood默认为"开心"，Date默认为当天日期。')

ap(doc, '核心代码——数据模型：')
ac(doc, 'public class DiaryEntry')
ac(doc, '{')
ac(doc, '    public string Id { get; set; } = Guid.NewGuid().ToString("N")[..8];')
ac(doc, '    public string Title { get; set; } = "";')
ac(doc, '    public string Content { get; set; } = "";')
ac(doc, '    public string Mood { get; set; } = "开心";')
ac(doc, '    public DateTime Date { get; set; } = DateTime.Today;')
ac(doc, '    public DateTime CreatedAt { get; set; } = DateTime.Now;')
ac(doc, '}')

ap(doc,
    '（2）数据持久化模块。使用System.Text.Json将日记列表序列化为JSON格式保存到本地diaries.json文件。'
    '程序启动时通过LoadData方法自动加载已有数据，每次修改后通过SaveData方法写回文件。'
    '设置WriteIndented = true使JSON格式化输出便于人工查看和调试。'
    '两个方法均采用try-catch异常处理，确保文件读写失败时程序不会崩溃，'
    '且在加载失败时自动初始化为空列表。')

ap(doc, '核心代码——数据持久化：')
ac(doc, 'private void LoadData()')
ac(doc, '{')
ac(doc, '    try { if (File.Exists(_dataFile))')
ac(doc, '        _entries = JsonSerializer.Deserialize<List<DiaryEntry>>(File.ReadAllText(_dataFile)) ?? new(); }')
ac(doc, '    catch { _entries = new(); }')
ac(doc, '}')
ac(doc, '')
ac(doc, 'private void SaveData()')
ac(doc, '{')
ac(doc, '    try')
ac(doc, '    {')
ac(doc, '        File.WriteAllText(_dataFile, JsonSerializer.Serialize(_entries,')
ac(doc, '            new JsonSerializerOptions { WriteIndented = true }));')
ac(doc, '    }')
ac(doc, '    catch (Exception ex) { MessageBox.Show("保存失败：" + ex.Message); }')
ac(doc, '}')

ap(doc,
    '（3）日记列表展示模块。使用ListBox的OwnerDraw模式实现自定义绘制列表项，每行紧凑显示日期（8pt字体）、'
    '标题（10pt加粗，超长截断加省略号）、心情标签三个信息。通过MeasureItem设置每项高度为44像素以容纳两行文本，'
    '偶数行和奇数行使用不同背景色（深色交替），选中行使用主题色（#EB9150）高亮。'
    '数据按日记日期倒序排列，确保最新日记始终显示在最前面。')

ap(doc,
    '（4）日记编辑模块。右侧编辑面板包含标题文本框、正文多行文本框、日期选择器（DateTimePicker）和心情下拉框'
    '（ComboBox，提供开心、一般、难过、兴奋、疲惫五种心情选项）。通过布尔标志_isEditing区分新建和编辑两种模式：'
    '新建时清空表单并自动聚焦标题输入框，编辑时将选中日记的已有数据加载到各控件中。'
    '保存时对标题进行非空校验，确保每条日记都有有效标题。删除操作弹出确认对话框，防止误删。')

ap(doc,
    '（5）搜索模块。在左侧搜索框输入关键字时，通过TextChanged事件实时触发列表刷新。'
    '使用LINQ的Where方法结合string.Contains按标题或正文内容进行忽略大小写的模糊匹配，'
    '关键字为空或空白字符时显示全部日记。搜索结果同样保持日期倒序排列。')

ap(doc, '3.2 AI辅助编程实践分析', bold=True)

ap(doc,
    '在DailyNote项目的开发过程中，多次使用Claude Code辅助编码。以下从三个典型功能点的AI辅助编程实践出发，'
    '分析AI在Windows程序设计中的优势与局限。')

ap(doc,
    '（1）UI自定义绘制——AI提供框架思路，人工修正API细节。'
    '在ListBox自定义绘制功能的实现中，向Claude Code提供的Prompt为："使用WinForm的ListBox实现日记列表'
    '自定义绘制，每行显示日期、标题、心情标签三个信息，要求偶数行和奇数行背景色不同，选中行高亮显示。"'
    'AI正确生成了OwnerDraw模式的整体框架，包括设置DrawMode为OwnerDrawVariable、在MeasureItem中设置'
    'ItemHeight为44像素、在DrawItem中通过Graphics对象绘制多行文本。但人工审核发现AI使用了不存在的'
    'e.Item属性（DrawItemEventArgs中无此属性），需通过ListBox.Items[e.Index]获取数据项。'
    '此外AI使用的固定颜色值与项目的深色主题不匹配，需要手动调整。'
    '分析结论：AI在提供实现思路和框架代码方面高效可靠，但在具体的API使用细节上存在"幻觉"现象，'
    '需要开发者具备基本的.NET框架知识进行验证和修正。')

ap(doc,
    '（2）数据持久化——标准化API调用准确率高。'
    '在数据持久化功能的实现中，Prompt为："使用System.Text.Json将日记列表序列化为JSON格式保存到本地'
    'diaries.json文件中，程序启动时自动加载已有数据，要求包含异常处理。"'
    'AI生成的LoadData和SaveData方法几乎无需修改即可直接使用，仅在异常提示语言（英译中）和catch块中'
    '的空列表初始化方面做了少量调整。对于System.Text.Json这类API设计规范、用法固定的标准化库，'
    'AI生成代码的准确率很高。'
    '分析结论：AI最适合处理API标准化、逻辑固定的编码任务，这类任务的输出质量接近可直接投产的水平。')

ap(doc,
    '（3）搜索功能——核心逻辑正确，边界条件需人工补充。'
    '在搜索功能的实现中，Prompt为："在日记本中添加搜索功能，用户在搜索框输入关键字时实时过滤日记列表，'
    '按标题或正文内容匹配，忽略大小写，关键字为空时显示全部。"'
    'AI正确生成了基于TextChanged事件的实时过滤逻辑，使用LINQ的Where和string.Contains进行模糊匹配。'
    '但需要人工补充两个关键细节：一是空关键字或空白字符时应显示全部而非进行空字符串匹配（AI遗漏了边界判断），'
    '二是搜索结果应保持按日期倒序排列（AI默认未添加排序逻辑）。'
    '分析结论：AI能快速生成核心业务逻辑，但对边界条件和异常场景的处理需要人工逐项检查和补充完善。')

ap(doc,
    'AI辅助编程实践总结。综合以上三个案例的实践经验，可以得出以下结论：'
    '第一，AI在快速生成标准化代码框架、提供常见API用法参考方面具有显著效率优势，'
    '能够大幅减少重复性编码工作，使开发者可以专注于核心逻辑和设计决策。'
    '第二，AI输出存在三个需要警惕的主要问题：一是API细节可能不准确（访问不存在的属性或方法），'
    '二是边界条件和异常场景的处理不完善（空值、边界值等），'
    '三是缺乏对项目整体风格一致性的考虑（颜色主题、命名规范等）。'
    '第三，AI辅助编程的最佳实践模式为"AI生成框架+人工审核把关"——先由AI快速搭建代码框架和实现标准逻辑，'
    '再由开发者逐行审查代码的正确性、完整性和风格一致性。开发者需要对AI输出的每一行代码负责，'
    '在享受效率提升的同时保持批判性判断，不能盲目信任和直接使用AI生成的代码。')

ap(doc, '3.3 运行结果截图', bold=True)

aimg(doc, r'D:\C#\期末作业\DailyNote\screenshot.png', 13)

ap(doc,
    '以上为DailyNote日记本应用程序的运行界面截图。截图展示了左侧日记列表（含搜索框、写日记按钮、'
    '日记项自定义绘制列表）和右侧编辑面板（含标题、正文、日期选择器、心情下拉框及保存/删除/清空按钮）的'
    '完整布局。界面采用暖色调（米白背景#F8F3EB搭配深色侧边栏#32374B），'
    '体现了Windows Forms控件编程、GDI+自定义绘制和数据持久化技术的综合应用成果。')

# ======================= 4. 项目总结反思 =======================
ap(doc, '4. 项目总结反思（30分）', bold=True, indent=False, size=Pt(14))
ap(doc, '4.1 课程知识综合应用总结（10分）', bold=True)

ap(doc,
    '在本学期的期末项目开发中，我完成了MusicPlayer音乐播放器、ZooMonitor动物园健康监测系统和DailyNote日记本'
    '三个不同领域的Windows应用程序。在这些项目的开发过程中，综合应用了本学期所学的多项课程知识。')

ap(doc,
    '第一，C#面向对象编程知识的应用。在DailyNote项目中定义了DiaryEntry数据模型类，包含Id、Title、Content、'
    'Mood、Date、CreatedAt六个属性，使用泛型List<DiaryEntry>管理日记集合，使用Guid生成唯一标识。'
    '在MusicPlayer项目中通过枚举类型MediaState管理播放状态，通过类封装播放控制逻辑。这些实践加深了对'
    '面向对象编程中类、对象、封装、泛型等核心概念的理解。')

ap(doc,
    '第二，WinForm控件编程的综合运用。三个项目共使用了Panel、Button、Label、TextBox、DataGridView、'
    'ListView、ListBox、DateTimePicker、ComboBox、TrackBar、Timer等十余种Windows Forms控件。'
    '通过为每个控件精细配置属性（Size、Location、BackColor、ForeColor、Font等）和绑定事件处理程序，'
    '实现了各具特色的用户界面和交互逻辑。特别是在ZooMonitor项目中通过DataGridView的条件格式化和'
    'SelectionChanged事件联动实现了动态数据展示。')

ap(doc,
    '第三，多媒体处理技术的应用。在MusicPlayer中使用axWindowsMediaPlayer控件实现音频播放控制，'
    '这是本学期课程重点内容之一。通过调用Ctlcontrols.play()/pause()/stop()方法和settings.volume'
    '属性，实现了完整的播放控制功能。')

ap(doc,
    '第四，文件操作与数据持久化技术的应用。在DailyNote项目中，通过File.ReadAllText和File.WriteAllText'
    '结合System.Text.Json序列化技术实现了数据的读取与存储。在MusicPlayer中使用File.Exists检测文件'
    '存在性，使用Path类处理文件路径。同时在这些文件操作的关键环节都添加了try-catch异常处理机制。')

ap(doc,
    '第五，GDI+绘图技术的应用。在ZooMonitor项目中，使用Graphics对象在Panel上绘制健康指标柱状图，'
    '使用LinearGradientBrush绘制标题栏渐变背景。在MusicPlayer中使用Graphics绘制CD光盘装饰图案和'
    '播放按钮圆形背景。')

ap(doc,
    '项目开发中遇到的主要问题：一是ListBox和ListView的OwnerDraw自定义绘制模式在API使用上与传统方式有差异，'
    '通过查阅微软官方文档和反复调试解决了绘制闪烁和文字对齐问题。二是System.Text.Json在处理中文内容时的'
    '编码问题，通过设置JsonSerializerOptions确保了中文字符的正确序列化。三是不同项目中控件样式的一致性'
    '维护，通过统一定义颜色常量和复用辅助方法提高了代码的可维护性。')

ap(doc,
    '收获与不足：通过本学期的课程学习和项目实践，我对C# Windows程序设计的完整开发流程有了系统的认识，'
    '从需求分析、UI设计、编码实现到测试调试，每个环节都积累了宝贵的实践经验。不足之处在于对设计模式'
    '的应用还不够熟练，代码架构的合理性有待提升，后续需要进一步学习和实践。')

ap(doc, '4.2 项目开发总结与反思（10分）', bold=True)

ap(doc,
    '项目开发过程回顾：三个项目的开发均采用RDD（需求驱动开发）模式。首先明确功能需求列表，然后设计UI界面布局，'
    '再进行编码实现，最后进行功能测试和细节优化。MusicPlayer侧重多媒体播放控制技术，DailyNote侧重数据管理'
    '和界面交互，ZooMonitor侧重数据可视化和仪表盘设计，三个项目各有侧重，共实践了本学期课程的大部分核心技术。')

ap(doc,
    'Claude Code在本次开发中提供了以下有效帮助：在UI布局方面，快速生成了界面框架代码和控件配置，节省了大量'
    '重复性编码时间。在功能逻辑方面，提供了数据序列化、Linq查询等标准化操作的代码参考。在技术方案方面，'
    '提供了OwnerDraw自定义绘制、条件格式化等实现思路和技术选型建议。')

ap(doc,
    '需要人工重点把关的方面：一是API调用的准确性，AI有时会生成不存在的属性或方法，需要开发者对.NET框架API'
    '有基本的了解才能发现和修正。二是边界条件处理，AI容易忽略空值、边界值和异常场景的处理。三是代码与项目'
    '整体风格的一致性，需要对AI生成的代码进行调整以匹配项目的命名规范和编码风格。')

ap(doc,
    '通过本次项目实践，我深刻体会到：AI辅助开发工具能够显著提升编码效率，但它本质上是一个"高级代码补全工具"，'
    '无法替代开发者对业务逻辑的理解、对代码质量的把控和对程序架构的设计。正确的使用方式是让AI承担重复性、'
    '标准化的编码工作，而将核心设计决策、逻辑验证和质量管控掌握在自己手中。')

ap(doc, '4.3 学习规划与发展方向（10分）', bold=True)

ap(doc,
    '对照当前就业市场对Windows客户端开发工程师的岗位要求（参考字节跳动校园招聘Windows客户端开发工程师岗位，'
    '工作地点成都），我认真分析了自身知识体系的优势与不足，并制定了针对性的学习计划。')

ap(doc,
    '岗位要求分析：该岗位要求熟练掌握C/C++、C#、Go、Rust中的至少一门语言，具备扎实的数据结构和算法基础，'
    '了解操作系统相关核心原理。参考链接：https://jobs.bytedance.com/campus/position/7530950785564231944/detail')

ap(doc,
    '自身不足：一是对操作系统底层原理（如内存管理、进程线程调度、文件系统）的了解还不够深入，需要在理论学习'
    '上补强。二是数据结构和算法能力需要系统训练，目前仅停留在课程教学的基础水平，距离企业面试要求有较大差距。'
    '三是缺乏大型项目的开发经验，对软件工程规范、设计模式应用和代码重构等方面理解不够。')

ap(doc,
    '短期学习计划（1-2个月）：深入学习C#高级特性，包括LINQ高级查询、异步编程（async/await）、反射技术、'
    '特性（Attribute）编程等；完善当前项目的功能和代码质量，作为求职作品展示。')

ap(doc,
    '中期学习计划（3-6个月）：系统学习数据结构和算法，通过LeetCode平台每周完成3-5道算法题，重点突破数组、'
    '链表、树、图、动态规划、贪心算法等核心题型。同时阅读《CLR via C#》和《深入理解Windows操作系统》'
    '两本经典技术书籍，提升对.NET运行时和Windows系统原理的理解。')

ap(doc,
    '长期学习计划（6-12个月）：拓展WPF和UWP桌面开发技术栈，学习MVVM设计模式。争取通过学校实习渠道进入'
    '企业参与实际项目开发，积累工程实践经验。同时持续跟进AI辅助开发工具的最新发展，在保证代码质量的前提下'
    '不断提升开发效率。')

# ======================= 5. 致谢 =======================
ap(doc, '5. 致谢', bold=True, indent=False, size=Pt(14))
ap(doc,
    '在本学期的课程学习和期末项目开发过程中，首先要感谢肖老师的悉心教导和耐心讲解。肖老师的课堂教学内容充实、'
    '条理清晰，既有理论知识的系统讲解，又有实践操作的详细演示，使我对Windows程序设计这门课程有了全面的认识'
    '和浓厚的学习兴趣。在项目开发过程中遇到的多个技术难题，也是通过回顾课堂笔记和课程实验代码找到的解决方案。')
ap(doc,
    '感谢Claude Code在项目开发中提供的代码辅助，在UI布局代码生成、技术方案参考、代码调试等方面提供了有效帮助。')
ap(doc,
    '感谢同学们在学习过程中的交流与讨论，特别是在技术问题上的互相帮助和启发，使我在遇到困难时能够更快找到解决思路。')
ap(doc,
    '最后感谢家人一直以来的支持和鼓励，为我提供了良好的学习环境和条件。')

# ======================= 6. 参考文献 =======================
ap(doc, '6. 参考文献', bold=True, indent=False, size=Pt(14))
ap(doc, '[1] 微软官方文档. Windows Forms 控件[EB/OL]. https://learn.microsoft.com/zh-cn/dotnet/desktop/winforms/')
ap(doc, '[2] 微软官方文档. AxWindowsMediaPlayer 控件[EB/OL]. https://learn.microsoft.com/zh-cn/windows/win32/wmp/')
ap(doc, '[3] 字节跳动校园招聘. Windows客户端开发工程师[EB/OL]. https://jobs.bytedance.com/campus/position/7530950785564231944/detail')
ap(doc, '[4] 微软官方文档. System.Text.Json 概述[EB/OL]. https://learn.microsoft.com/zh-cn/dotnet/standard/serialization/system-text-json/')
ap(doc, '[5] 微软官方文档. DataGridView 类[EB/OL]. https://learn.microsoft.com/zh-cn/dotnet/api/system.windows.forms.datagridview')

# ========== Save ==========
out = r'D:\C#\期末作业\Windows程序设计课程作业.docx'
doc.save(out)

total = sum(len(p.text) for p in doc.paragraphs if p.text.strip())
code_chars = 0
for p in doc.paragraphs:
    t = p.text
    if t.strip() and p.runs and p.runs[0].font.name == 'Courier New':
        code_chars += len(t)

# Also check for image paragraphs
img_count = 0
for p in doc.paragraphs:
    for r in p.runs:
        if r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
            img_count += 1
            break

print(f'Saved! Total chars: {total}')
print(f'Code chars (Courier New): {code_chars}')
print(f'Image paragraphs: {img_count}')
print(f'Non-code text: {total - code_chars}')
